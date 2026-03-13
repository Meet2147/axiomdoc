from __future__ import annotations

from pathlib import Path

from axiomdoc.models import Block, CanonicalDocument
from axiomdoc.parsers.base import ParserBackend


class DocxParser(ParserBackend):
    name = "python-docx"
    supported_suffixes = (".docx",)

    def parse(self, path: Path) -> CanonicalDocument:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError(
                "DOCX parsing requires the 'python-docx' dependency. Install 'axiomdoc[docx]' or 'axiomdoc[full]'."
            ) from exc

        source = Document(str(path))
        document = CanonicalDocument.empty(path, source_format="docx", parser_name=self.name)

        core = source.core_properties
        for field_name in ("title", "author", "subject", "keywords", "category", "comments"):
            value = getattr(core, field_name, None)
            if value:
                document.metadata[field_name] = value

        block_counter = 1

        for paragraph in source.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = paragraph.style.name if paragraph.style is not None else ""
            kind, level = self._classify_paragraph(style_name, text)
            document.blocks.append(
                Block(
                    block_id=f"d{block_counter}",
                    kind=kind,
                    text=text,
                    level=level,
                    metadata={"style": style_name or "Normal"},
                )
            )
            block_counter += 1

        for table_index, table in enumerate(source.tables, start=1):
            rows: list[list[str]] = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    rows.append(cells)
            if not rows:
                continue

            document.blocks.append(
                Block(
                    block_id=f"t{table_index}",
                    kind="table",
                    text="\n".join(" | ".join(row) for row in rows),
                    metadata={
                        "table_index": table_index,
                        "table_rows": rows,
                        "column_count": max(len(row) for row in rows),
                    },
                )
            )

        return document

    @staticmethod
    def _classify_paragraph(style_name: str, text: str) -> tuple[str, int | None]:
        normalized = style_name.strip().lower()
        if normalized.startswith("heading"):
            suffix = normalized.replace("heading", "", 1).strip()
            try:
                level = max(1, min(int(suffix or "1"), 6))
            except ValueError:
                level = 1
            return "heading", level
        if normalized.startswith("title"):
            return "heading", 1
        if normalized.startswith("subtitle"):
            return "heading", 2
        if normalized.startswith("list"):
            return "list_item", None
        return "paragraph", None
