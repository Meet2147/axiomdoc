from __future__ import annotations

from pathlib import Path

import fitz
from docx import Document
from openpyxl import Workbook

from axiomdoc.parsers.docx import DocxParser
from axiomdoc.parsers.pdf import PdfParser
from axiomdoc.parsers.registry import ParserRegistry
from axiomdoc.parsers.xlsx import XlsxParser
from axiomdoc.parsers.xml import XmlParser


def test_parser_registry_resolves_expected_backends() -> None:
    registry = ParserRegistry()

    assert isinstance(registry.resolve(Path("sample.pdf")), PdfParser)
    assert isinstance(registry.resolve(Path("sample.docx")), DocxParser)
    assert isinstance(registry.resolve(Path("sample.xlsx")), XlsxParser)
    assert isinstance(registry.resolve(Path("sample.xml")), XmlParser)


def test_pdf_parser_smoke(tmp_path: Path) -> None:
    pdf_path = tmp_path / "smoke.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Annual Report", fontsize=20)
    page.insert_text((72, 108), "Revenue increased by 20 percent.", fontsize=12)
    doc.save(pdf_path)
    doc.close()

    parsed = PdfParser().parse(pdf_path)

    assert parsed.source_format == "pdf"
    assert any(block.kind == "heading" for block in parsed.blocks)
    assert any("Revenue increased" in block.text for block in parsed.blocks)


def test_docx_parser_extracts_headings_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("Project Overview", level=1)
    document.add_paragraph("A short summary paragraph.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Field"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Status"
    table.cell(1, 1).text = "Active"
    document.save(path)

    parsed = DocxParser().parse(path)

    assert parsed.source_format == "docx"
    assert parsed.blocks[0].kind == "heading"
    assert any(block.kind == "table" for block in parsed.blocks)


def test_xlsx_parser_extracts_sheet_headings_and_rows(tmp_path: Path) -> None:
    path = tmp_path / "sample.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Revenue"
    sheet["A1"] = "Region"
    sheet["B1"] = "Revenue"
    sheet["A2"] = "North"
    sheet["B2"] = 120
    workbook.save(path)
    workbook.close()

    parsed = XlsxParser().parse(path)

    assert parsed.source_format == "xlsx"
    assert parsed.blocks[0].kind == "heading"
    assert parsed.blocks[0].text == "Revenue"
    assert any(block.kind == "table" for block in parsed.blocks[1:])
