from __future__ import annotations

from pathlib import Path

import fitz
from docx import Document
from openpyxl import Workbook
from PIL import Image, ImageDraw

from axiomdoc.exporters.markdown import document_to_markdown
from axiomdoc.exporters.xml import document_to_xml
from axiomdoc.indexing import build_index_manifest
from axiomdoc.parsers.docx import DocxParser
from axiomdoc.parsers.pdf import PdfParser
from benchmarks.labeled_eval import evaluate_fixture


def test_table_blocks_render_to_markdown_and_xml_and_index(tmp_path: Path) -> None:
    path = tmp_path / "table.docx"
    document = Document()
    document.add_heading("Project Overview", level=1)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Field"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Status"
    table.cell(1, 1).text = "Active"
    document.save(path)

    parsed = DocxParser().parse(path)
    markdown = document_to_markdown(parsed)
    xml_output = document_to_xml(parsed)
    index_manifest = build_index_manifest(parsed)

    assert "| Field | Value |" in markdown
    assert "<table><row><cell>Field</cell><cell>Value</cell></row>" in xml_output
    assert any(chunk["metadata"]["kind"] == "table" for chunk in index_manifest["chunks"])


def test_pdf_parser_ocr_fallback_on_image_only_page(tmp_path: Path) -> None:
    pdf_path = tmp_path / "ocr.pdf"
    png_path = tmp_path / "ocr.png"

    image = Image.new("RGB", (800, 200), color="white")
    draw = ImageDraw.Draw(image)
    draw.text((40, 60), "Scanned OCR Text", fill="black")
    image.save(png_path)

    pdf = fitz.open()
    page = pdf.new_page()
    rect = fitz.Rect(50, 50, 350, 150)
    page.insert_image(rect, filename=str(png_path))
    pdf.save(pdf_path)
    pdf.close()

    parsed = PdfParser().parse(pdf_path)

    assert any("Scanned OCR Text" in block.text for block in parsed.blocks)


def test_labeled_evaluation_scores_expected_headings_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "labeled.docx"
    document = Document()
    document.add_heading("Project Overview", level=1)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Field"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Status"
    table.cell(1, 1).text = "Active"
    document.save(path)

    labels = Path("tests/fixtures/sample_labels.json")
    result = evaluate_fixture(path, labels)

    assert result["summary"]["heading_precision"] == 1.0
    assert result["summary"]["heading_recall"] == 1.0
    assert result["summary"]["table_precision"] == 1.0
    assert result["summary"]["table_recall"] == 1.0


def test_xlsx_table_rows_become_multiple_index_chunks(tmp_path: Path) -> None:
    from axiomdoc.parsers.xlsx import XlsxParser

    path = tmp_path / "sheet.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Metrics"
    sheet.append(["Quarter", "Revenue"])
    sheet.append(["Q1", 120])
    sheet.append(["Q2", 135])
    workbook.save(path)
    workbook.close()

    parsed = XlsxParser().parse(path)
    manifest = build_index_manifest(parsed)

    table_chunks = [chunk for chunk in manifest["chunks"] if chunk["metadata"]["kind"] == "table"]
    assert len(table_chunks) == 3
